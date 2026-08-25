"""app/services/docx_builder.py(4단계, 데이터 조회 없는 순수 함수) 단위 테스트.

조회 로직이 전혀 없어서 고정 fixture만으로 표/차트/빈 섹션 처리를 검증할 수 있다.
"""

from __future__ import annotations

import io
from datetime import date

from docx import Document

from app.services.docx_builder import SectionResult, _render_chart, build_report_docx


def _reopen(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))


def _headings(document: Document) -> list[str]:
    return [
        p.text
        for p in document.paragraphs
        if p.style.name == "Title" or p.style.name.startswith("Heading")
    ]


def test_build_report_docx_returns_reopenable_bytes_with_cover_and_section_headings() -> (
    None
):
    sections = [
        SectionResult(title="매출 추이", narrative="월별 매출입니다."),
    ]

    result = build_report_docx(
        "월간 매출 리포트", date(2026, 1, 1), date(2026, 3, 31), sections
    )

    assert isinstance(result, bytes)
    assert len(result) > 0
    document = _reopen(result)
    assert _headings(document) == ["월간 매출 리포트", "매출 추이"]


def test_build_report_docx_cover_page_states_the_query_period() -> None:
    result = build_report_docx("리포트", date(2026, 1, 1), date(2026, 3, 31), [])
    document = _reopen(result)
    body_text = "\n".join(p.text for p in document.paragraphs)
    assert "2026-01-01 ~ 2026-03-31" in body_text


def test_build_report_docx_creates_one_table_per_section_with_rows() -> None:
    sections = [
        SectionResult(
            title="매출 추이",
            narrative="월별 매출입니다.",
            columns=["order_month", "total_sales"],
            rows=[["2026-01", 1000], ["2026-02", 1200]],
        ),
        SectionResult(
            title="고객별 순위",
            narrative="고객 순위입니다.",
            columns=["customer_name", "total_sales"],
            rows=[["Acme", 500]],
        ),
    ]

    result = build_report_docx(
        "월간 매출 리포트", date(2026, 1, 1), date(2026, 3, 31), sections
    )
    document = _reopen(result)

    assert len(document.tables) == 2
    first_table = document.tables[0]
    assert [cell.text for cell in first_table.rows[0].cells] == [
        "order_month",
        "total_sales",
    ]
    assert [cell.text for cell in first_table.rows[1].cells] == ["2026-01", "1000"]
    assert len(first_table.rows) == 3  # 헤더 1 + 데이터 행 2


def test_build_report_docx_skips_table_and_chart_for_empty_section() -> None:
    sections = [
        SectionResult(title="미수금 현황", narrative="해당 기간 데이터가 없습니다.")
    ]

    result = build_report_docx(
        "월간 매출 리포트", date(2026, 1, 1), date(2026, 3, 31), sections
    )
    document = _reopen(result)

    assert len(document.tables) == 0
    assert len(document.inline_shapes) == 0
    assert "해당 기간 데이터가 없습니다." in [p.text for p in document.paragraphs]


def test_build_report_docx_adds_chart_only_when_chartable_and_chart_type_set() -> None:
    chartable_section = SectionResult(
        title="매출 추이",
        narrative="월별 매출입니다.",
        columns=["order_month", "total_sales"],
        rows=[["2026-01", 1000], ["2026-02", 1200]],
        chartable=True,
        chart_type="line",
        label_column="order_month",
        value_column="total_sales",
    )
    # chartable=False면 label/value 컬럼이 있어도 차트를 만들지 않는다.
    non_chartable_section = SectionResult(
        title="고객별 순위",
        narrative="고객 순위입니다.",
        columns=["customer_name", "total_sales"],
        rows=[["Acme", 500]],
        chartable=False,
        chart_type="bar",
        label_column="customer_name",
        value_column="total_sales",
    )

    result = build_report_docx(
        "월간 매출 리포트",
        date(2026, 1, 1),
        date(2026, 3, 31),
        [chartable_section, non_chartable_section],
    )
    document = _reopen(result)

    assert len(document.tables) == 2
    assert len(document.inline_shapes) == 1


def test_render_chart_returns_none_when_label_column_not_in_columns() -> None:
    section = SectionResult(
        title="깨진 섹션",
        narrative="",
        columns=["order_month", "total_sales"],
        rows=[["2026-01", 1000]],
        chartable=True,
        chart_type="bar",
        label_column="missing_column",
        value_column="total_sales",
    )

    assert _render_chart(section) is None
