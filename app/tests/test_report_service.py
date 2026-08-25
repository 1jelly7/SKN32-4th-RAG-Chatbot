"""app/services/report_service.py(5단계) 오케스트레이션 단위 테스트.

FakeMCPPort/질문을 보고 응답을 고르는 fake port로 실제 네트워크·DB 없이 섹션별
병렬 조회 → build_tables 변환 → docx 조립까지의 흐름을 검증한다.
"""

from __future__ import annotations

import asyncio
import io
from datetime import date
from typing import Any

import pytest
from docx import Document

from app.mcp.client import FakeMCPPort, MCPClient, MCPForbiddenError
from app.services.report_service import generate_report
from app.services.report_templates import get_template

TEST_USER = {"role": "admin"}


def _block(
    rows: list[dict[str, Any]],
    sql: str = "SELECT 1",
    label: str = "",
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    return {
        "label": label,
        "generated_sql": sql,
        "rows": rows,
        "row_count": len(rows),
        "metadata": metadata or {},
    }


def _success(domain: str, data: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "status": "success",
        "domain": domain,
        "message": None,
        "data": data,
        "sources": [],
        "metadata": {},
    }


def _no_result(domain: str) -> dict[str, Any]:
    return {
        "status": "error",
        "domain": domain,
        "message": "해당 기간의 데이터가 없습니다.",
        "error_code": "NO_RESULT",
        "data": [],
        "sources": [],
        "metadata": {},
    }


class _QuestionAwarePort:
    """질문 문구로 어느 섹션 호출인지 구분해 서로 다른 응답을 돌려주는 fake."""

    def __init__(self, responses: dict[str, dict[str, Any]]) -> None:
        self._responses = responses
        self.questions: list[str] = []

    async def call_tool(self, tool_name: str, payload: dict[str, Any]) -> object:
        question = str(payload.get("question", ""))
        self.questions.append(question)
        for keyword, response in self._responses.items():
            if keyword in question:
                return response
        raise AssertionError(f"예상하지 못한 질문: {question}")


def test_generate_report_fills_period_placeholders_into_each_section_question() -> None:
    port = _QuestionAwarePort(
        {
            "추이": _success("sales", [_block([{"order_month": "2026-01"}])]),
            "순위": _success("sales", [_block([{"customer_name": "Acme"}])]),
            "미수금": _success("sales", [_block([{"outstanding": 1}])]),
        }
    )

    asyncio.run(
        generate_report(
            "sales_monthly",
            date(2026, 1, 1),
            date(2026, 3, 31),
            MCPClient(port),
            user_context=TEST_USER,
        )
    )

    assert len(port.questions) == 3
    for question in port.questions:
        assert "2026-01-01" in question
        assert "2026-03-31" in question


def test_generate_report_produces_a_table_per_section_with_data() -> None:
    port = _QuestionAwarePort(
        {
            "추이": _success(
                "sales",
                [
                    _block(
                        [{"order_month": "2026-01", "total_sales": 1000}],
                        metadata={"chart_hint": "line"},
                    )
                ],
            ),
            "순위": _success(
                "sales",
                [
                    _block(
                        [{"customer_name": "Acme", "total_sales": 500}],
                        metadata={"chart_hint": "bar"},
                    )
                ],
            ),
            "미수금": _success(
                "sales", [_block([{"customer_name": "Acme", "outstanding": 10}])]
            ),
        }
    )

    docx_bytes = asyncio.run(
        generate_report(
            "sales_monthly",
            date(2026, 1, 1),
            date(2026, 3, 31),
            MCPClient(port),
            user_context=TEST_USER,
        )
    )

    document = Document(io.BytesIO(docx_bytes))
    assert len(document.tables) == 3
    # 매출 추이/고객별 순위는 chart_hint가 있어 차트가 붙고, 미수금은 없다.
    assert len(document.inline_shapes) == 2


def test_generate_report_uses_empty_placeholder_narrative_for_no_result_section() -> (
    None
):
    port = _QuestionAwarePort(
        {
            "추이": _success(
                "sales", [_block([{"order_month": "2026-01", "total_sales": 1000}])]
            ),
            "순위": _success(
                "sales", [_block([{"customer_name": "Acme", "total_sales": 500}])]
            ),
            "미수금": _no_result("sales"),
        }
    )

    docx_bytes = asyncio.run(
        generate_report(
            "sales_monthly",
            date(2026, 1, 1),
            date(2026, 3, 31),
            MCPClient(port),
            user_context=TEST_USER,
        )
    )

    document = Document(io.BytesIO(docx_bytes))
    assert len(document.tables) == 2  # 미수금 섹션은 표가 없다
    paragraphs = [p.text for p in document.paragraphs]
    assert any("데이터가 없습니다" in text for text in paragraphs)


def test_generate_report_propagates_non_no_result_mcp_errors() -> None:
    class _ForbiddenPort:
        async def call_tool(self, tool_name: str, payload: dict[str, Any]) -> object:
            return {
                "status": "error",
                "domain": "sales",
                "message": "권한 없음",
                "error_code": "FORBIDDEN",
                "data": [],
                "sources": [],
                "metadata": {},
            }

    with pytest.raises(MCPForbiddenError):
        asyncio.run(
            generate_report(
                "sales_monthly",
                date(2026, 1, 1),
                date(2026, 3, 31),
                MCPClient(_ForbiddenPort()),
                user_context=TEST_USER,
            )
        )


def test_generate_report_raises_key_error_for_unknown_template() -> None:
    with pytest.raises(KeyError):
        asyncio.run(
            generate_report(
                "no-such-template",
                date(2026, 1, 1),
                date(2026, 3, 31),
                MCPClient(FakeMCPPort({})),
                user_context=TEST_USER,
            )
        )


def test_sales_monthly_template_sections_stay_in_sales_domain() -> None:
    """report_service가 도메인별로 옳은 MCP tool을 고르는지는, 이 템플릿의 모든
    섹션이 sales라는 전제에 기대고 있다 — 그 전제가 깨지면 이 테스트가 먼저 잡는다.
    """
    template = get_template("sales_monthly")
    assert {section["domain"] for section in template["sections"]} == {"sales"}
    assert len(template["sections"]) == 3
