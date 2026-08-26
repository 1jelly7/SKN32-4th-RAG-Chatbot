"""python-docx/matplotlib 설치 확인용 최소 스모크 테스트(2단계).

리포트 기능(docx_builder, report_service)이 아직 없는 시점에도, 두 의존성이
설치돼 있고 기본 동작(문서 생성/재오픈, 차트 PNG 렌더링)을 한다는 것만 확인한다.
"""

from __future__ import annotations

import io

import matplotlib

matplotlib.use("Agg")  # 헤드리스 환경에서 디스플레이 없이 렌더링하기 위함.

import matplotlib.pyplot as plt
from docx import Document


def test_python_docx_creates_and_reopens_document() -> None:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("스모크 테스트", level=1)
    document.add_paragraph("python-docx 설치 확인용 문단입니다.")
    document.save(buffer)

    buffer.seek(0)
    reopened = Document(buffer)
    paragraphs = [p.text for p in reopened.paragraphs]
    assert "스모크 테스트" in paragraphs
    assert "python-docx 설치 확인용 문단입니다." in paragraphs


def test_matplotlib_renders_chart_to_png_bytes() -> None:
    figure, axis = plt.subplots()
    axis.bar(["A", "B", "C"], [1, 2, 3])

    buffer = io.BytesIO()
    figure.savefig(buffer, format="png")
    plt.close(figure)

    png_bytes = buffer.getvalue()
    assert png_bytes.startswith(b"\x89PNG\r\n\x1a\n")
    assert len(png_bytes) > 0
