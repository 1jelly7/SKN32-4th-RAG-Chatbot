"""판매/구매 조회 결과를 .docx 리포트로 조립하는 순수 함수.

이 모듈은 데이터 조회를 전혀 하지 않는다 — MCP 호출, Text2SQL, evidence 평가는
전부 report_service.py(5단계)의 책임이고, 여기서는 이미 만들어진 SectionResult만
받아 문서 bytes로 직렬화한다. 조회 로직이 없어서 고정 fixture로 바로 단위테스트할
수 있다(app/tests/test_docx_builder.py).
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date
from typing import Any, Literal

import matplotlib

matplotlib.use("Agg")  # 서버 환경에는 디스플레이가 없어, 렌더링에 GUI 백엔드가 필요 없게 한다.

import matplotlib.pyplot as plt
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from matplotlib import font_manager

# 추가: 차트 라벨(월·고객명 등)이 전부 한글인데, matplotlib 기본 폰트(DejaVu Sans)는
# 한글 글리프가 없어 빈 네모(tofu)로 렌더링된다. 배포 환경마다 설치된 폰트가 달라
# 하나로 하드코딩하지 않고, 설치된 후보 중 첫 번째를 찾아 쓴다 — 하나도 없으면
# 조용히 기본값을 둔다(차트 자체는 여전히 만들어지고, 라벨만 깨진다).
_KOREAN_FONT_CANDIDATES = (
    "Malgun Gothic",
    "NanumGothic",
    "Noto Sans KR",
    "Noto Sans CJK KR",
    "AppleGothic",
)


def _configure_korean_font() -> None:
    installed = {f.name for f in font_manager.fontManager.ttflist}
    for candidate in _KOREAN_FONT_CANDIDATES:
        if candidate in installed:
            plt.rcParams["font.family"] = candidate
            break
    plt.rcParams["axes.unicode_minus"] = False


_configure_korean_font()

# 추가: 위 _KOREAN_FONT_CANDIDATES는 "서버에 설치된" 폰트 중에서 골라 matplotlib
# 차트 PNG를 그릴 때만 쓴다(이미지는 서버에서 미리 렌더링되므로). 반면 표지·문단·표
# 셀 같은 docx 네이티브 텍스트는 서버가 아니라 뷰어(대부분 Windows의 Word)가 열 때
# 그려지므로, "서버에 뭐가 깔려 있는지"가 아니라 "뷰어에서 어떤 폰트를 쓸지"가
# 중요하다. python-docx 기본 스타일은 동아시아(eastAsia) 폰트를 지정하지 않아서,
# Word는 보통 알아서 대체하지만 그렇지 않은 뷰어는 한글을 빈 네모로 그린다. 그래서
# Windows 기본 한글 UI 폰트를 문서에 직접 못박아둔다.
_DOCX_KOREAN_FONT = "맑은 고딕"


def _set_run_korean_font(run: Any, font_name: str = _DOCX_KOREAN_FONT) -> None:
    """run의 서양(ascii)·동아시아(eastAsia) 폰트를 모두 지정한다.

    run.font.name만 설정하면 w:rFonts의 ascii만 채워지고 eastAsia는 비어 있는
    채로 남는다. eastAsia를 직접 채워야 한글이 폰트 대체 없이 그대로 남지 않는다.
    """
    run.font.name = font_name
    run_properties = run._element.get_or_add_rPr()
    fonts_element = run_properties.find(qn("w:rFonts"))
    if fonts_element is None:
        fonts_element = OxmlElement("w:rFonts")
        run_properties.append(fonts_element)
    fonts_element.set(qn("w:eastAsia"), font_name)


def _apply_korean_font(document: DocxDocument) -> None:
    """표지·섹션 헤딩·문단·표 셀 전체의 모든 run에 한글 폰트를 적용한다."""
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            _set_run_korean_font(run)


@dataclass(frozen=True)
class SectionResult:
    """리포트 한 섹션. app/agent/nodes.py::build_tables()가 만드는 표 항목 1개에
    제목·서술문단을 붙인 형태라, report_service.py(5단계)가 그 결과를 거의 그대로
    옮겨 담을 수 있게 필드를 맞췄다.
    """

    title: str
    narrative: str
    columns: list[str] = field(default_factory=list)
    rows: list[list[Any]] = field(default_factory=list)
    chartable: bool = False
    chart_type: Literal["bar", "line"] | None = None
    label_column: str | None = None
    value_column: str | None = None


def build_report_docx(
    title: str,
    start_date: date,
    end_date: date,
    sections: list[SectionResult],
) -> bytes:
    """표지 뒤에 섹션을 순서대로 담은 .docx를 만들어 bytes로 반환한다."""
    document = Document()
    _add_cover_page(document, title, start_date, end_date)
    for section in sections:
        _add_section(document, section)
    _apply_korean_font(document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_cover_page(
    document: DocxDocument, title: str, start_date: date, end_date: date
) -> None:
    document.add_heading(title, level=0)
    document.add_paragraph(f"조회 기간: {start_date.isoformat()} ~ {end_date.isoformat()}")
    document.add_page_break()


def _add_section(document: DocxDocument, section: SectionResult) -> None:
    document.add_heading(section.title, level=1)
    document.add_paragraph(section.narrative)

    if not section.columns or not section.rows:
        return

    _add_table(document, section.columns, section.rows)

    if section.chartable and section.chart_type and section.label_column and section.value_column:
        chart_png = _render_chart(section)
        if chart_png is not None:
            document.add_picture(io.BytesIO(chart_png), width=Inches(5.5))


def _add_table(
    document: DocxDocument, columns: list[str], rows: list[list[Any]]
) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for index, name in enumerate(columns):
        header_cells[index].text = str(name)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)


def _render_chart(section: SectionResult) -> bytes | None:
    """label_column/value_column 조합으로 막대·꺾은선 차트를 PNG bytes로 그린다.

    컬럼 이름이 실제 columns 목록과 어긋나면(방어적으로만) 차트를 건너뛴다 — 표는
    이미 추가됐으므로 차트 하나 때문에 문서 전체 생성이 실패하지 않게 한다.
    """
    try:
        label_index = section.columns.index(section.label_column)
        value_index = section.columns.index(section.value_column)
    except ValueError:
        return None

    labels = [str(row[label_index]) for row in section.rows]
    values = [row[value_index] for row in section.rows]

    figure, axis = plt.subplots(figsize=(6, 3.5))
    if section.chart_type == "line":
        axis.plot(labels, values, marker="o")
    else:
        axis.bar(labels, values)
    axis.set_title(section.title)
    figure.autofmt_xdate(rotation=45)

    buffer = io.BytesIO()
    figure.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(figure)
    return buffer.getvalue()
